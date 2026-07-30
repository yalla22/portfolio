# -*- coding: utf-8 -*-
"""Generate role-tailored resumes for Janadri Yalla Yashwanth from Marut Drones projects."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ---------- shared person data ----------
NAME = "Janadri Yalla Yashwanth"
EMAIL = "yallayashwanth99@gmail.com"
PHONE = "+91 8790819924"
LOCATION = "Hyderabad, India 500001"
LINKEDIN_URL = "https://www.linkedin.com/in/janadri-yalla-yashwanth-9579b5306/"
GITHUB_URL = "https://github.com/yalla22"
MARUT_URL = "https://www.linkedin.com/company/marutdrones"

EDUCATION = [
    ("Sri Venkateswara College of Engineering, Tirupati", "Dec 2021 – May 2025",
     "B.Tech in Computer Science (Artificial Intelligence & Machine Learning)  —  CGPA: 8.6"),
    ("Narayana Junior College, Hyderabad", "Jun 2019 – Mar 2021",
     "Intermediate (MPC)  —  96%"),
]

CERTS = [
    "CS50’s Introduction to Programming with Python — Harvard University Online",
    "Databases: Topics in SQL — Stanford Online",
]

ACADEMIC_ML = [
    ("Classification of Sleep Disorders", "Python, TensorFlow",
     "Built an ML model classifying sleep-disorder types from physiological/behavioral data; cleaned and preprocessed data and achieved 98% classification accuracy, served via a web app for early diagnosis."),
    ("Heart Disease Prediction", "Python, scikit-learn",
     "Trained ML models on clinical features (age, BP, cholesterol, heart rate) to predict heart-disease risk, delivering a low-cost, scalable preventive-screening tool."),
    ("Twitter Sentiment Analysis", "Python, NLP, TF-IDF",
     "Classified tweets into positive/negative/neutral sentiment using NLP preprocessing (tokenization, stop-word removal) and TF-IDF feature extraction for customer-feedback insight."),
]

# ---------- project library ----------
P = {
    "flightlog": {
        "title": "Flight Log Analyzer — Full-Stack Drone Diagnostics Platform",
        "tech": "React, TypeScript, FastAPI, Python, Docker",
        "bullets": [
            "Built a full-stack web app that ingests raw drone flight logs and auto-generates professional PDF health/crash reports, supporting both ArduPilot (.bin) and JIYI K++ formats with automatic binary-vs-text detection via magic-byte sniffing.",
            "Engineered a from-scratch binary parser for the ArduPilot/DataFlash self-describing log format using Python struct, decoding 50+ message types and reconstructing real-world IST timestamps from GPS-week data.",
            "Implemented a phased diagnostic engine running 50+ automated checks (power, IMU/vibration, GPS, EKF, compass, RC, failsafes), classifying findings as PASS/WARN/FAIL and inferring probable crash cause.",
            "Developed automated PDF reporting with ReportLab + Matplotlib (flight-path maps, sensor charts, color-coded tables); built the React 19 + Vite + Tailwind front end with file upload and download handling.",
            "Containerized the full stack with Docker / Docker Compose; integrated a .NET converter via Mono + Xvfb inside Linux (~7,200 LOC, solo-built).",
        ],
    },
    "geoai": {
        "title": "GeoAI Detection — Buildings & Trees from Drone Orthomosaics",
        "tech": "Python, PyTorch, LangSAM (SAM + GroundingDINO), DeepForest, rasterio, OpenCV",
        "bullets": [
            "Built a geospatial building-detection pipeline on drone orthomosaics using zero-shot LangSAM (SAM + GroundingDINO) with a “building roof” text prompt, processing large GeoTIFFs block-wise via rasterio (no training data required).",
            "Eliminated flat false positives (roads, bare fields) by filtering every detection against the DSM — computing height = roof − ground (nDSM) and dropping objects without real vertical structure; exported georeferenced GeoJSON with area and height.",
            "Built individual-tree-crown detection with DeepForest (RetinaNet, NEON-pretrained), GSD-aware resampling tuned for dense tropical canopies, HSV green-filtering, crown merging, and NMS.",
            "Separated real trees from grass using a Canopy Height Model (CHM = DSM − DTM), keeping only crowns ≥ 1.5 m tall — removing flat green vegetation that color filters alone cannot reject.",
        ],
    },
    "aerial": {
        "title": "Aerial Image Object Detection — Multi-Model Benchmark",
        "tech": "Python, PyTorch, Hugging Face Transformers, SegFormer, OpenCV",
        "bullets": [
            "Developed a computer-vision pipeline detecting buildings, vehicles, and trees from high-resolution drone imagery (stitched panoramas up to 14,400 × 7,200 px) using SegFormer-B5 (ADE20K, 150 classes) with PyTorch and Hugging Face Transformers.",
            "Benchmarked 4+ segmentation models (SegFormer-B5/B2, Aerial-Drone SegFormer, SAM + CLIP) head-to-head and selected ADE20K for best precision, cutting false positives ~3× versus over-segmenting alternatives.",
            "Engineered the full inference pipeline in OpenCV/NumPy: sliding-window tiling, HSV sky-skipping, semantic-mask-to-instance contour extraction, geometric shape filtering, and per-class Non-Max Suppression.",
            "Delivered annotated outputs (per-class count overlays + JSON), batch-folder processing, and drone-video support (frame-skip → annotated MP4), running on GPU/CUDA at ~7–11 s per image.",
        ],
    },
    "precland": {
        "title": "Autonomous Precision Landing on Moving Platform",
        "tech": "ArduPilot, MAVLink, Raspberry Pi 5, AprilTag, OpenCV, Python",
        "bullets": [
            "Built an autonomous drone precision-landing system that lands on an AprilTag target on stationary and moving (ground-rover) platforms; achieved proven stationary-tag landings at ~5–10 cm accuracy.",
            "Wrote the Raspberry Pi 5 companion-computer controller (pupil-apriltags + picamera2 + OpenCV) streaming LANDING_TARGET and velocity setpoints to a CubeOrange (ArduCopter) over MAVLink (pymavlink) at 20 Hz.",
            "Designed a three-mode state machine — visual P-control with rover-velocity feed-forward, GPS-NAV haversine approach, and hover fallback — with pilot RC override, velocity caps, and touchdown auto-disarm.",
            "Implemented multi-vehicle MAVLink (drone + rover SYSIDs), camera calibration, systemd auto-start, and a BENCH safety mode for props-off bench testing.",
        ],
    },
    "survey": {
        "title": "Marut Survey Platform — Backend & AI Detection Services",
        "tech": "FastAPI, Celery, Redis, PostgreSQL, PyTorch, SAM-2, S3/GDAL, Docker",
        "bullets": [
            "Built three production detection pipelines (stockpile, tree-crown, building) integrating three model families — SAM-2, DeepForest, and GroundingDINO/LangSAM — behind a FastAPI backend with async Celery/Redis job APIs (live progress, cancellation, owner-scoped RBAC).",
            "Derived per-object stockpile volumes (m³) and material classification from nDSM height integration, backed by a 7-factor confidence rubric and Union-Find cross-tile NMS dedup at 0.3 IoU.",
            "Streamed cloud-optimized GeoTIFFs block-by-block from S3 via GDAL /vsis3/ + rasterio (no full-raster load) and tuned S3 multipart transfers — 64 MB chunks, 25 concurrent threads, pooled connections — for multi-gigabyte orthomosaics.",
            "Packaged the PyTorch / SAM-2 / DeepForest / GroundingDINO stack into a Docker image shipped through GitHub Actions to systemd-managed GPU and CPU workers; Pydantic v2 schemas and Alembic migrations.",
        ],
    },
}

# ---------- role configs ----------
ROLES = {
    "ML_CV_Engineer": {
        "filename": "Resume_Janadri_Yalla_Yashwanth_ML_CV_Engineer",
        "headline": "Computer Vision / Machine Learning Engineer",
        "summary": "Computer-vision and ML engineer with hands-on production experience building deep-learning detection pipelines on drone and aerial imagery — SegFormer, SAM-2, DeepForest, and GroundingDINO/LangSAM — paired with geospatial post-processing (rasterio/GDAL, DSM/DTM elevation analysis). B.Tech in CSE (AI & ML), CGPA 8.6. Strong in PyTorch, OpenCV, and shipping models behind real APIs.",
        "skills": [
            ("Languages", "Python, C, SQL"),
            ("ML / Deep Learning", "PyTorch, Hugging Face Transformers, TensorFlow, scikit-learn, CUDA/GPU"),
            ("Computer Vision", "SegFormer, SAM / SAM-2, DeepForest, GroundingDINO / LangSAM, CLIP, OpenCV, NumPy"),
            ("Geospatial", "rasterio, GDAL, GeoTIFF / COG, DSM/DTM & CHM, GeoJSON, CRS / projections"),
            ("Backend / MLOps", "FastAPI, Celery, Redis, Docker, GitHub Actions, REST APIs"),
            ("Concepts", "Semantic segmentation, object detection, NMS, model benchmarking, data preprocessing"),
        ],
        "tools": "Git & GitHub, Claude Code, VS Code, Jupyter, Hugging Face Hub, Docker, GitHub Actions, QGIS, Linux",
        "projects": ["aerial", "geoai", "survey", "flightlog", "precland"],
        "academic": [],
    },
    "Backend_Python_Engineer": {
        "filename": "Resume_Janadri_Yalla_Yashwanth_Backend_Python_Engineer",
        "headline": "Backend / Python Engineer",
        "summary": "Backend engineer building production Python services with FastAPI — asynchronous job orchestration (Celery + Redis), cloud data streaming from S3 (GDAL/rasterio), role-based access control, and containerized deployment with Docker and CI/CD. Comfortable owning a feature end-to-end from API design to GPU/CPU worker deployment. B.Tech in CSE (AI & ML), CGPA 8.6.",
        "skills": [
            ("Languages", "Python, C, SQL, TypeScript"),
            ("Backend / Frameworks", "FastAPI, Uvicorn (ASGI), SQLAlchemy 2.0, Pydantic v2, REST API design, async/await"),
            ("Async / Data", "Celery, Redis, PostgreSQL / PostGIS, S3 / MinIO, Alembic migrations, GDAL, rasterio"),
            ("Auth / Security", "JWT, RBAC, CORS, file-upload validation"),
            ("DevOps", "Docker, Docker Compose, GitHub Actions CI/CD, systemd, Linux deployment"),
            ("ML integration", "PyTorch, SAM-2, DeepForest, OpenCV (serving CV models behind APIs)"),
        ],
        "tools": "Git & GitHub, Claude Code, VS Code, Docker, GitHub Actions, Postman, Linux, systemd",
        "projects": ["survey", "flightlog", "aerial", "geoai"],
        "academic": [],
    },
    "Backend_Developer": {
        "filename": "Resume_Janadri_Yalla_Yashwanth_Backend_Developer",
        "headline": "Backend / Python Engineer",
        "summary": "Backend engineer who ships production Python services end to end. At Marut Drones I build FastAPI systems backed by Celery, Redis, and PostgreSQL, stream large geospatial rasters from S3 with GDAL and rasterio, and serve computer-vision models behind async job APIs. I own features from API design through Docker packaging, CI/CD, and Linux deployment. B.Tech in CS (AI & ML), CGPA 8.6; available immediately.",
        "skills": [
            ("Languages", "Python, SQL, JavaScript / TypeScript, C"),
            ("Backend & APIs", "FastAPI, Uvicorn (ASGI), REST APIs, async/await, Pydantic v2, SQLAlchemy 2.0"),
            ("Async & Data", "Celery, Redis, PostgreSQL / PostGIS, Alembic, S3 (AWS) / MinIO"),
            ("Geospatial", "GDAL, rasterio, GeoTIFF / COG, GeoJSON, CRS / projections"),
            ("Auth & Security", "JWT, RBAC, CORS, request validation"),
            ("DevOps", "Docker, Docker Compose, GitHub Actions, CI/CD, systemd, Linux"),
            ("ML integration", "PyTorch, OpenCV, SAM-2, DeepForest (served behind APIs)"),
        ],
        "tools": "Git & GitHub, Claude Code, VS Code, Postman, Docker, GitHub Actions, Linux",
        "marut_bullets": [
            "Built three production geospatial detection pipelines integrating three model families — SAM-2, DeepForest, and GroundingDINO/LangSAM — behind a FastAPI backend, with async job APIs that enqueue Celery tasks on Redis and stream live progress and cancellation.",
            "Derived per-object stockpile volumes (m³) and material classification from nDSM height integration, backed by a 7-factor confidence rubric and Union-Find cross-tile NMS dedup at 0.3 IoU.",
            "Streamed orthomosaics block-by-block from S3 via GDAL /vsis3/ (512 MB cache, GDAL_NUM_THREADS=ALL_CPUS), eliminating full-raster memory buffering for multi-gigabyte COGs.",
            "Tuned S3 multipart transfers (64 MB chunks, 25 concurrent threads, pooled boto3 connections) and moved Celery workers to adaptive concurrency (nproc-2) for faster large-file I/O.",
            "Implemented owner-scoped RBAC and JWT auth with Pydantic v2 validation and Alembic migrations; packaged the PyTorch/SAM-2 stack into a Docker image shipped via GitHub Actions to systemd GPU/CPU workers.",
        ],
        "rooman_bullet": "Developed and evaluated a sleep-disorder classification model on physiological and behavioral data, reaching 98% accuracy, and delivered it through a web app for clinicians.",
        "projects_heading": "Selected Projects",
        "custom_projects": [
            {
                "title": "Flight Log Analyzer — Drone Diagnostics Backend",
                "tech": "FastAPI, Python, asyncio, ReportLab, Docker",
                "bullets": [
                    "Built a FastAPI service that ingests raw drone flight logs and returns generated PDF diagnostics, with async uploads, 500 MB limits, extension allow-listing, and format auto-detection by magic-byte sniffing.",
                    "Wrote a binary parser from scratch for the ArduPilot/DataFlash format using Python struct, decoding 50+ message types and reconstructing real wall-clock timestamps from GPS-week data.",
                    "Serialized a headless .NET converter under Mono and Xvfb with an asyncio lock, per-job temp directories, and background cleanup, then rendered diagnostics with ReportLab and Matplotlib.",
                    "Containerized the service with Docker and Docker Compose for reproducible deployment (~7,200 lines, solo-built).",
                ],
            },
            {
                "title": "Aerial Image Object Detection — Multi-Model Benchmark",
                "tech": "Python, PyTorch, Hugging Face Transformers, SegFormer, OpenCV",
                "bullets": [
                    "Benchmarked 4+ segmentation models (SegFormer-B5/B2, an aerial-domain model, SAM+CLIP) on identical drone imagery and selected SegFormer-B5 (ADE20K) for the cleanest detections, cutting false positives roughly 3x versus over-segmenting baselines.",
                    "Built the inference pipeline in OpenCV/NumPy — sliding-window tiling on panoramas up to 14,400x7,200 px, HSV sky-skipping, mask-to-instance contour extraction, and per-class NMS — emitting annotated images plus structured JSON.",
                ],
            },
        ],
        "projects": [],
        "academic": [],
    },
    "Backend_Developer_Java_Mongo": {
        "filename": "Resume_Janadri_Yalla_Yashwanth_Backend_Developer_Java_MongoDB",
        "headline": "Backend Engineer | Python · Java · FastAPI · SQL & NoSQL",
        "summary": "Backend engineer who ships production Python services end to end, with working knowledge of Java and both SQL and NoSQL data stores. At Marut Drones I build FastAPI systems backed by Celery, Redis, and PostgreSQL, stream large geospatial rasters from S3 with GDAL and rasterio, and serve computer-vision models behind async job APIs. I own features from API design through Docker packaging, CI/CD, and Linux deployment. B.Tech in CS (AI & ML), CGPA 8.6; available immediately.",
        "skills": [
            ("Languages", "Python, Java, SQL, JavaScript / TypeScript, C"),
            ("Backend & APIs", "FastAPI, Uvicorn (ASGI), REST APIs, async/await, Pydantic v2, SQLAlchemy 2.0"),
            ("Databases", "PostgreSQL / PostGIS, MongoDB, Redis, Alembic"),
            ("Async & Cloud", "Celery, Redis, S3 (AWS) / MinIO"),
            ("Geospatial", "GDAL, rasterio, GeoTIFF / COG, GeoJSON, CRS / projections"),
            ("Auth & Security", "JWT, RBAC, CORS, request validation"),
            ("DevOps", "Docker, Docker Compose, GitHub Actions, CI/CD, systemd, Linux"),
            ("ML integration", "PyTorch, OpenCV, SAM-2, DeepForest (served behind APIs)"),
        ],
        "tools": "Git & GitHub, Claude Code, VS Code, Postman, Docker, GitHub Actions, Linux",
        "marut_bullets": [
            "Built three production geospatial detection pipelines integrating three model families — SAM-2, DeepForest, and GroundingDINO/LangSAM — behind a FastAPI backend, with async job APIs that enqueue Celery tasks on Redis and stream live progress and cancellation.",
            "Derived per-object stockpile volumes (m³) and material classification from nDSM height integration, backed by a 7-factor confidence rubric and Union-Find cross-tile NMS dedup at 0.3 IoU.",
            "Streamed orthomosaics block-by-block from S3 via GDAL /vsis3/ (512 MB cache, GDAL_NUM_THREADS=ALL_CPUS), eliminating full-raster memory buffering for multi-gigabyte COGs.",
            "Tuned S3 multipart transfers (64 MB chunks, 25 concurrent threads, pooled boto3 connections) and moved Celery workers to adaptive concurrency (nproc-2) for faster large-file I/O.",
            "Implemented owner-scoped RBAC and JWT auth with Pydantic v2 validation and Alembic migrations; packaged the PyTorch/SAM-2 stack into a Docker image shipped via GitHub Actions to systemd GPU/CPU workers.",
        ],
        "rooman_bullet": "Developed and evaluated a sleep-disorder classification model on physiological and behavioral data, reaching 98% accuracy, and delivered it through a web app for clinicians.",
        "projects_heading": "Selected Projects",
        "custom_projects": [
            {
                "title": "Flight Log Analyzer — Drone Diagnostics Backend",
                "tech": "FastAPI, Python, asyncio, ReportLab, Docker",
                "bullets": [
                    "Built a FastAPI service that ingests raw drone flight logs and returns generated PDF diagnostics, with async uploads, 500 MB limits, extension allow-listing, and format auto-detection by magic-byte sniffing.",
                    "Wrote a binary parser from scratch for the ArduPilot/DataFlash format using Python struct, decoding 50+ message types and reconstructing real wall-clock timestamps from GPS-week data.",
                    "Containerized the service with Docker and Docker Compose for reproducible deployment (~7,200 lines, solo-built).",
                ],
            },
            {
                "title": "Aerial Image Object Detection — Multi-Model Benchmark",
                "tech": "Python, PyTorch, Hugging Face Transformers, SegFormer, OpenCV",
                "bullets": [
                    "Benchmarked 4+ segmentation models (SegFormer-B5/B2, an aerial-domain model, SAM+CLIP) on identical drone imagery and selected SegFormer-B5 (ADE20K) for the cleanest detections, cutting false positives roughly 3x versus over-segmenting baselines.",
                    "Built the inference pipeline in OpenCV/NumPy — sliding-window tiling on panoramas up to 14,400x7,200 px, HSV sky-skipping, mask-to-instance contour extraction, and per-class NMS.",
                ],
            },
        ],
        "projects": [],
        "academic": [],
    },
    "MLOps_Engineer": {
        "filename": "Resume_Janadri_Yalla_Yashwanth_MLOps_Engineer",
        "headline": "Machine Learning Engineer | MLOps · Model Training · Docker · CI/CD",
        "summary": "Machine learning engineer and 2025 CS (AI & ML) graduate who trains deep-learning models and ships them to production. I train custom object detectors (YOLO) on self-labeled data and deploy production detection models (SAM-2, DeepForest, SegFormer) behind FastAPI, owning the path from model to Docker image, GitHub Actions CI/CD, and systemd GPU/CPU workers. Comfortable across the ML lifecycle — data pipelines, training, validation, containerization, and deployment. B.Tech in CS (AI & ML), CGPA 8.6; immediate joiner.",
        "skills": [
            ("Languages", "Python, SQL, C, Java"),
            ("ML & Deep Learning", "PyTorch, TensorFlow, scikit-learn, YOLO (Ultralytics), SAM-2, DeepForest, SegFormer, OpenCV"),
            ("ML Lifecycle / MLOps", "Model training & validation, custom dataset labeling & augmentation, data pipelines, model packaging & serving"),
            ("Deployment & CI/CD", "Docker, Docker Compose, GitHub Actions, CI/CD, systemd, Linux"),
            ("Cloud & Data", "AWS S3, GDAL / rasterio, PostgreSQL, Redis, Celery"),
            ("Algorithms", "Classification, regression, clustering, object detection, deep learning, NMS"),
            ("Practices", "Git version control, testing & debugging, Agile collaboration"),
        ],
        "tools": "Git & GitHub, Claude Code, VS Code, Jupyter, Hugging Face Hub, Docker, GitHub Actions, Linux",
        "marut_bullets": [
            "Productionized three ML detection pipelines (SAM-2, DeepForest, GroundingDINO/LangSAM) behind a FastAPI backend with async Celery/Redis job orchestration, taking models from prototype to scalable serving.",
            "Owned the ML deployment path end to end — packaged the PyTorch model stack into reproducible Docker images and shipped them through GitHub Actions CI/CD to systemd-managed GPU and CPU workers.",
            "Engineered data pipelines that stream cloud-optimized rasters from S3 via GDAL /vsis3/ and rasterio, with elevation-based post-processing and cross-tile NMS deduplication for reliable model outputs.",
            "Debugged production model failures (CUDA/PyTorch mismatches, GPU-versus-CPU serving, failed migrations) to keep inference stable and reduce cycle time.",
        ],
        "rooman_bullet": "Trained and evaluated a sleep-disorder classification model (TensorFlow) on physiological and behavioral data, reaching 98% accuracy, and delivered it through a web app.",
        "projects_heading": "Selected Projects",
        "custom_projects": [
            {
                "title": "Custom Object Detection with YOLO",
                "tech": "Python, YOLO (Ultralytics), PyTorch, OpenCV, custom datasets",
                "bullets": [
                    "Trained and validated custom YOLO detectors on self-labeled datasets for number-plate, vehicle, water-body, and broken-fence detection.",
                    "Owned the full training pipeline — data collection, annotation, augmentation, train/validation splits, and per-task tuning.",
                ],
            },
            {
                "title": "Aerial Image Object Detection — Multi-Model Benchmark",
                "tech": "Python, PyTorch, Hugging Face Transformers, SegFormer, OpenCV",
                "bullets": [
                    "Benchmarked 4+ segmentation models on identical drone imagery and selected SegFormer-B5 (ADE20K) for the cleanest detections, cutting false positives roughly 3x versus baselines.",
                    "Built the OpenCV/NumPy inference pipeline — tiling, filtering, mask-to-instance extraction, and per-class NMS — with annotated image and JSON outputs.",
                ],
            },
        ],
        "projects": [],
        "academic": [],
    },
    "CPP_Developer": {
        "filename": "Resume_Janadri_Yalla_Yashwanth_CPP_Developer",
        "headline": "C++ Developer | OOP · Data Structures & Algorithms · Problem Solving",
        "summary": "2025 B.Tech CS graduate with a strong programming foundation in C++, C, and Python, and solid command of OOP and data structures & algorithms. Currently a software engineer at Marut Drones, where I write clean, modular code, debug production issues, and collaborate in a fast-moving team. Strong logical and analytical thinking, and eager to grow as a C++ developer. No backlogs; immediate joiner, open to relocation.",
        "skills": [
            ("Languages", "C++, C, Python, Java, SQL"),
            ("C++ & OOP", "Classes & objects, inheritance, polymorphism, encapsulation, STL, pointers & memory management, references"),
            ("CS Fundamentals", "Data Structures & Algorithms, Operating Systems, DBMS, problem solving"),
            ("Software Practices", "Clean code, debugging, testing, Git version control, Agile collaboration"),
            ("Also worked with", "FastAPI, React, Docker, Linux, OpenCV, PyTorch"),
        ],
        "tools": "Git & GitHub, VS Code, GCC / g++, Claude Code, Docker, Linux",
        "marut_bullets": [
            "Software engineer on the AI/automation team building production systems — writing clean, modular code, applying OOP design, debugging production issues, and collaborating on design and delivery.",
            "Built backend services and data pipelines (FastAPI, Celery/Redis) with version control, code review, and testing best practices.",
            "Diagnosed and resolved production failures (environment/config mismatches, deployment issues) to keep systems stable.",
        ],
        "rooman_bullet": "Built and evaluated a machine-learning classification model on healthcare data, reaching 98% accuracy.",
        "projects_heading": "Projects",
        "custom_projects": [
            {
                "title": "Flight Log Analyzer — Binary Log Parser & Diagnostics",
                "tech": "Binary parsing (byte-level struct), Python, Docker",
                "bullets": [
                    "Wrote a binary parser from scratch for a self-describing log format, decoding 50+ message types with byte-level parsing, bit fields, and timestamp reconstruction.",
                    "Implemented a 50+ check diagnostic engine and packaged the app with Docker (~7,200 lines, solo-built).",
                ],
            },
            {
                "title": "Autonomous Precision Landing — Real-Time Control",
                "tech": "Real-time control loops, state machines, MAVLink, Python",
                "bullets": [
                    "Designed a real-time control loop and multi-mode state machine (visual control, GPS navigation, hover fallback) running at 20 Hz for autonomous drone landing.",
                    "Handled coordinate geometry, velocity limits, and safety fallbacks, achieving ~5-10 cm landing accuracy.",
                ],
            },
        ],
        "projects": [],
        "academic": [],
    },
    "AI_Software_Engineer": {
        "filename": "Resume_Janadri_Yalla_Yashwanth_AI_Software_Engineer",
        "headline": "Software Engineer | Backend + AI Systems | Python · FastAPI · PyTorch · Model Deployment",
        "summary": "Software engineer and 2025 CS (AI & ML) graduate, CGPA 8.6, with a strong foundation in Data Structures & Algorithms, Operating Systems, DBMS, and OOP. At Marut Drones I build backend systems that deploy and serve ML models at scale — FastAPI microservices with Celery/Redis job orchestration, Dockerized GPU/CPU inference, and CI/CD. I write clean, scalable code and I'm eager to build LLM and GenAI systems. Immediate joiner.",
        "skills": [
            ("Languages", "Python, Java, C++, C, SQL"),
            ("CS Fundamentals", "Data Structures & Algorithms, Operating Systems, Computer Networks, DBMS, OOP"),
            ("Backend & Systems", "FastAPI, REST APIs, microservices, async/await, Celery, Redis, distributed task queues, Pydantic, SQLAlchemy"),
            ("Databases", "PostgreSQL, SQL, Redis"),
            ("AI / ML", "PyTorch, deep learning, model deployment & inference at scale, OpenCV, YOLO, SAM-2, DeepForest"),
            ("GenAI / LLM", "Prompt engineering (AI-assisted development with Claude Code); learning RAG, embeddings & AI agents"),
            ("Cloud / DevOps", "Docker, GitHub Actions, CI/CD, AWS S3, Git, Linux"),
        ],
        "tools": "Git & GitHub, Claude Code, VS Code, Jupyter, Docker, GitHub Actions, Postman, Linux",
        "marut_bullets": [
            "Built and productionized ML detection services (SAM-2, DeepForest, GroundingDINO/LangSAM) behind a FastAPI backend, serving models at scale with async Celery/Redis job orchestration (live progress, cancellation, RBAC).",
            "Owned model deployment and inference infrastructure — packaged the PyTorch stack into Docker images shipped via GitHub Actions to systemd-managed GPU and CPU workers.",
            "Engineered data pipelines streaming multi-gigabyte rasters block-by-block from S3 (GDAL/rasterio), with CRS-aware processing, elevation math, and cross-tile deduplication.",
            "Debugged production inference issues (CUDA/PyTorch mismatches, GPU-versus-CPU serving, failed migrations) and applied clean-code, testing, and version-control practices across the team.",
        ],
        "rooman_bullet": "Trained and evaluated a machine-learning classification model (TensorFlow) on healthcare data, reaching 98% accuracy.",
        "projects_heading": "Selected Projects",
        "custom_projects": [
            {
                "title": "Custom Object Detection with YOLO",
                "tech": "Python, YOLO (Ultralytics), PyTorch, OpenCV, custom datasets",
                "bullets": [
                    "Trained and validated custom YOLO detectors on self-labeled datasets for number-plate, vehicle, water-body, and broken-fence detection.",
                    "Owned the full pipeline — data collection, annotation, augmentation, train/validation splits, and per-task tuning.",
                ],
            },
            {
                "title": "Flight Log Analyzer — Full-Stack Diagnostics Backend",
                "tech": "FastAPI, Python, asyncio, React, Docker",
                "bullets": [
                    "Built a FastAPI service that ingests raw drone logs and returns generated PDF diagnostics, with a from-scratch binary parser decoding 50+ message types and a 50+ check diagnostic engine.",
                    "Containerized the full stack with Docker/Docker Compose for reproducible deployment (~7,200 lines, solo-built).",
                ],
            },
        ],
        "projects": [],
        "academic": [],
    },
    "Admin_Assistant": {
        "filename": "Resume_Janadri_Yalla_Yashwanth_Administrative_Assistant",
        "headline": "Administrative Support | MS Office · Documentation · Coordination",
        "summary": "Detail-oriented B.Tech graduate with hands-on experience producing accurate documentation and deliverables and coordinating across teams in a fast-paced environment. Proficient in Microsoft Office (Word, Excel, PowerPoint, Outlook), with strong organization, communication, and time-management skills. Comfortable managing multiple assignments to deadlines while maintaining quality standards. Open to rotational shifts; immediate joiner based in Hyderabad.",
        "skills": [
            ("Microsoft Office", "Word, Excel (formulas, tables), PowerPoint, Outlook"),
            ("Administrative", "Business correspondence, document formatting & editing, data entry & management, workflow / task tracking"),
            ("Coordination", "Cross-team and cross-location coordination, stakeholder support, scheduling"),
            ("Strengths", "Attention to detail, accuracy, organization, time management, communication"),
        ],
        "marut_bullets": [
            "Prepared, formatted, and maintained detailed documentation and deliverables, keeping them accurate and consistent in a fast-paced startup environment.",
            "Coordinated with cross-functional teams to gather requirements and deliver work to deadlines.",
            "Managed multiple assignments at once with strong attention to detail, organization, and follow-through.",
        ],
        "rooman_bullet": "Documented and delivered project work, collaborating with the team and communicating results clearly.",
        "projects": [],
        "custom_projects": [],
        "academic": [],
    },
    "Data_Scientist": {
        "filename": "Resume_Janadri_Yalla_Yashwanth_Data_Scientist",
        "headline": "Data Scientist | Machine Learning · Statistical Modeling · Python",
        "summary": "Data scientist and 2025 CS (AI & ML) graduate, CGPA 8.6, who applies machine learning, statistical modeling, and data analysis to real problems. At Marut Drones I train, evaluate, and deploy ML models on drone and aerial data, and I've built classification and prediction models across healthcare and NLP. Strong in Python (pandas, NumPy, scikit-learn, PyTorch) with a focus on turning data into actionable insights. Immediate joiner.",
        "skills": [
            ("Languages", "Python, SQL"),
            ("Machine Learning", "Classification, regression, clustering, deep learning, model training & evaluation"),
            ("ML Libraries", "scikit-learn, PyTorch, TensorFlow, OpenCV, YOLO"),
            ("Data Analysis", "pandas, NumPy, data cleaning, feature engineering, EDA, Matplotlib"),
            ("Statistics", "Statistical modeling, hypothesis-driven analysis, metrics (accuracy / precision / recall)"),
            ("Data & Deployment", "SQL, data pipelines, Jupyter, Git, Docker, FastAPI (model serving)"),
        ],
        "tools": "Git & GitHub, Jupyter, Claude Code, VS Code, Docker, Linux",
        "marut_bullets": [
            "Trained, validated, and deployed machine-learning models (object detection, segmentation) on drone and aerial imagery, owning the pipeline from data preparation to production serving.",
            "Built data pipelines that clean, transform, and analyze large geospatial datasets (imagery + elevation) to produce actionable outputs — areas, volumes, and counts.",
            "Applied model evaluation and error analysis and benchmarked multiple models to select the best performer for each detection task.",
            "Collaborated with cross-functional teams to turn raw drone data into business-relevant insights.",
        ],
        "rooman_bullet": "Built and evaluated a sleep-disorder classification model on physiological and behavioral data using Python and TensorFlow, achieving 98% accuracy after data cleaning and feature engineering.",
        "projects_heading": "Projects",
        "custom_projects": [
            {
                "title": "Heart Disease Prediction",
                "tech": "Python, scikit-learn, pandas",
                "bullets": [
                    "Trained classification models on clinical features (age, blood pressure, cholesterol, heart rate) to predict heart-disease risk, with data cleaning, feature engineering, and model evaluation.",
                ],
            },
            {
                "title": "Twitter Sentiment Analysis",
                "tech": "Python, NLP, TF-IDF, scikit-learn",
                "bullets": [
                    "Classified tweets into positive / negative / neutral using NLP preprocessing (tokenization, stop-word removal) and TF-IDF features; evaluated model performance for customer-feedback insight.",
                ],
            },
            {
                "title": "Custom Object Detection with YOLO",
                "tech": "Python, YOLO (Ultralytics), PyTorch, OpenCV",
                "bullets": [
                    "Trained custom YOLO detectors on self-labeled datasets (vehicle, number-plate, water, fence), owning data annotation, augmentation, and validation.",
                ],
            },
        ],
        "projects": [],
        "academic": [],
    },
    "SOC_Analyst": {
        "filename": "Resume_Janadri_Yalla_Yashwanth_SOC_Security_Analyst",
        "headline": "Security Operations (SOC) Analyst | Log & Event Analysis · Networking · Python · Linux",
        "summary": "Computer Science graduate (2025) moving into security operations, with a strong foundation in networking, Linux, and log/event analysis. I built an automated log-analysis and anomaly-investigation engine that triages events, separates false positives from real anomalies, and documents findings — the same analytical workflow behind SOC monitoring and incident triage. Strong Python and analytical skills; actively upskilling toward CompTIA Security+ and hands-on SOC labs. Willing to work 24/7 shifts; immediate joiner.",
        "skills": [
            ("Security & SOC", "Log & event analysis, event triage, false-positive filtering, incident investigation & documentation, phishing/email awareness; SIEM & incident response (actively learning)"),
            ("Networking", "TCP/IP, DNS, HTTP/S, OSI model, VPN & firewall concepts, network & log analysis"),
            ("Systems", "Linux, Windows, systemd, command line, service & system troubleshooting"),
            ("Scripting & Tools", "Python (automation & analysis), Bash, Git, Docker"),
            ("Fundamentals", "Operating Systems, Computer Networks, DBMS, Data Structures & Algorithms"),
            ("Strengths", "Analytical & investigative mindset, attention to detail, clear communication, calm under pressure"),
        ],
        "tools": "Git & GitHub, Linux, Python, VS Code, Wireshark & Splunk (learning), TryHackMe / LetsDefend labs",
        "marut_bullets": [
            "Software engineer operating production systems on Linux — analyzing logs, investigating failures, and tracing root causes across services, then documenting findings for the team.",
            "Built and maintained backend services and data pipelines with version control, monitoring, and incident-style troubleshooting under time pressure.",
            "Communicated technical issues and resolutions clearly to both technical and non-technical teammates.",
        ],
        "rooman_bullet": "Analyzed and modeled healthcare data in Python, applying a detail-oriented, evidence-based approach to reach reliable conclusions.",
        "projects_heading": "Projects",
        "custom_projects": [
            {
                "title": "Flight Log Analyzer — Automated Log Analysis & Anomaly Investigation",
                "tech": "Python, log parsing, anomaly detection, reporting",
                "bullets": [
                    "Built a platform that ingests raw logs, decodes them, and runs a 50+ check engine to triage events — separating false positives from genuine anomalies and investigating probable root causes.",
                    "Reverse-engineered an encrypted (XOR-obfuscated) binary log format to enable analysis, then generated structured investigation reports — mirroring SOC log analysis, triage, and incident documentation.",
                ],
            },
        ],
        "projects": [],
        "academic": [],
    },
    "Data_Auditing_Associate": {
        "filename": "Resume_Janadri_Yalla_Yashwanth_Data_Auditing_Associate",
        "headline": "Data Auditing & Annotation Associate | Image/Video Data · Accuracy · Attention to Detail",
        "summary": "B.Tech graduate with hands-on experience labeling and auditing image datasets for computer-vision models. Strong attention to detail, accuracy, and focus, and comfortable reviewing visual data on screen for extended periods to meet quality and productivity targets. Willing to work rotational and night shifts; remote-ready with a dedicated workspace. Immediate joiner.",
        "skills": [
            ("Data Auditing & Annotation", "Image/video data labeling, dataset annotation & validation, visual quality checking, accuracy & consistency"),
            ("Strengths", "Attention to detail, sustained focus, speed with accuracy, pattern recognition, sound judgement"),
            ("Work readiness", "Rotational & night shifts, remote/WFH-ready, meeting incremental targets, strong team player"),
            ("Computer skills", "MS Office, web-based tools, data entry, quick to learn new software"),
        ],
        "marut_bullets": [
            "Annotated and validated large image datasets for computer-vision models, labeling objects accurately and consistently across thousands of images.",
            "Reviewed model outputs and imagery for correctness, maintaining a high level of accuracy and flagging errors and edge cases.",
            "Worked to quality and productivity targets in a fast-paced environment, with a detail-oriented, consistent approach.",
        ],
        "rooman_bullet": "Cleaned, reviewed, and analyzed datasets with a careful, evidence-based approach to ensure accurate results.",
        "projects_heading": "Relevant Project",
        "custom_projects": [
            {
                "title": "Custom Dataset Labeling for Object Detection",
                "tech": "Image annotation, data validation, quality control",
                "bullets": [
                    "Annotated and validated self-labeled image datasets (vehicle, number-plate, water, fence) for object-detection models, ensuring labeling accuracy and consistency across the dataset.",
                ],
            },
        ],
        "projects": [],
        "academic": [],
    },
    "Java_Backend_Engineer": {
        "filename": "Resume_Janadri_Yalla_Yashwanth_Software_Engineer_Java_Backend",
        "headline": "Software Engineer | Backend & Microservices | Java · Python · REST · Docker · CI/CD",
        "summary": "Software engineer building production backend microservices — REST APIs, async job orchestration, and containerized deployment on Linux via CI/CD. Strong in OOP design and problem-solving, proficient in Python with working knowledge of Java, and actively expanding into Spring Boot and Kubernetes. B.Tech in CS (AI & ML), CGPA 8.6; comfortable in small, fast-paced, self-organizing teams.",
        "skills": [
            ("Languages", "Java, Python, SQL, C++, C"),
            ("Backend & Microservices", "REST APIs, microservices, FastAPI, OOP design, concurrency & async programming"),
            ("Databases", "PostgreSQL, Redis, SQL"),
            ("Containers & DevOps", "Docker, Docker Compose, CI/CD (GitHub Actions), Git, Linux / Unix"),
            ("Testing & Process", "Unit testing, SDLC, code review, debugging"),
            ("Networking & Fundamentals", "TCP/IP, HTTP, Computer Networks, Data Structures & Algorithms, Operating Systems"),
            ("Expanding into", "Spring Boot, Kubernetes, Kafka, Go (actively building proficiency)"),
        ],
        "tools": "Git & GitHub, Docker, Postman, IntelliJ IDEA, VS Code, Claude Code, Linux",
        "marut_bullets": [
            "Built production microservices with FastAPI and REST APIs, backed by Celery/Redis async job orchestration (live progress, cancellation, RBAC), applying OOP design and clean-code practices.",
            "Containerized services with Docker and shipped them through CI/CD (GitHub Actions) to Linux workers; debugged concurrency and production issues across distributed workers.",
            "Modeled and queried data in PostgreSQL and Redis, and designed and consumed REST APIs with schema validation.",
            "Collaborated in a small, fast-paced team and maintained clear technical documentation.",
        ],
        "rooman_bullet": "Built and evaluated a machine-learning model in Python, applying strong analytical and problem-solving skills.",
        "projects_heading": "Selected Projects",
        "custom_projects": [
            {
                "title": "Flight Log Analyzer — Full-Stack Diagnostics Backend",
                "tech": "Python, FastAPI, REST, asyncio, Docker",
                "bullets": [
                    "Built a FastAPI backend that ingests raw logs and returns generated reports, with a from-scratch binary parser (byte-level, 50+ message types) and a 50+ check diagnostic engine.",
                    "Containerized the full stack with Docker / Docker Compose for reproducible deployment (~7,200 lines, solo-built).",
                ],
            },
            {
                "title": "Marut Survey Platform — Detection Microservices",
                "tech": "FastAPI, Celery, Redis, PostgreSQL, Docker, S3",
                "bullets": [
                    "Built async detection microservices with Celery/Redis job queues, owner-scoped RBAC, and PostgreSQL, streaming large datasets from S3.",
                    "Shipped the stack via GitHub Actions CI/CD to systemd-managed workers on Linux.",
                ],
            },
        ],
        "projects": [],
        "academic": [],
    },
    "Software_Engineer_Python": {
        "filename": "Resume_Janadri_Yalla_Yashwanth_Software_Engineer_Python",
        "headline": "Software Engineer | Python · APIs · Clean Code · Agile",
        "summary": "Python software engineer and 2025 CSE graduate (CGPA 8.6) building production systems at Marut Drones — REST APIs, a full-stack web app, and ML-backed services. I write clean, maintainable Python, work in an agile team using Git and code review, and own features through testing, debugging, and deployment. Strong CS fundamentals; based in Hyderabad and available immediately.",
        "skills": [
            ("Languages", "Python, SQL, JavaScript / TypeScript, Java, C"),
            ("Python & Frameworks", "FastAPI, async/await, Pydantic, SQLAlchemy, REST APIs, NumPy, OpenCV"),
            ("Web / Full-Stack", "React, HTML, CSS, Vite"),
            ("Practices", "Agile collaboration, SDLC, code review, testing & debugging, clean & maintainable code"),
            ("Databases", "PostgreSQL, Redis, SQL"),
            ("CS Fundamentals", "Data Structures & Algorithms, DBMS, Operating Systems, OOP"),
            ("DevOps / Version Control", "Git & GitHub, GitHub Actions CI/CD, Docker, Linux"),
        ],
        "tools": "Git & GitHub, Claude Code, VS Code, Postman, Docker, GitHub Actions, Linux",
        "marut_bullets": [
            "Developed production Python services on the AI/automation team — three detection pipelines behind a FastAPI backend with async Celery/Redis job APIs — writing clean, modular code shared across detectors.",
            "Integrated SAM-2, DeepForest, and GroundingDINO/LangSAM, streaming large orthomosaics from S3 via GDAL/rasterio and deriving stockpile volumes (m³) and material classification.",
            "Debugged and stabilized production failures (CUDA/PyTorch mismatches, failed migrations, GPU/CPU worker setup), collaborating with the team on design and integration.",
            "Packaged the stack with Docker and shipped through GitHub Actions to systemd workers; used Git and code review for day-to-day version control.",
        ],
        "rooman_bullet": "Developed and evaluated a sleep-disorder classification model in Python, reaching 98% accuracy, and delivered it through a web app for clinicians.",
        "projects_heading": "Selected Projects",
        "custom_projects": [
            {
                "title": "Flight Log Analyzer — Full-Stack Drone Diagnostics",
                "tech": "Python, FastAPI, React, TypeScript, Docker",
                "bullets": [
                    "Built a full-stack app (FastAPI + React) that ingests raw drone logs and generates PDF diagnostic reports, with a from-scratch Python binary parser decoding 50+ message types.",
                    "Ran a 50+ check diagnostic engine and containerized the stack with Docker Compose for reproducible deployment (~7,200 lines, solo-built).",
                ],
            },
            {
                "title": "Aerial Image Object Detection — Multi-Model Benchmark",
                "tech": "Python, PyTorch, SegFormer, OpenCV",
                "bullets": [
                    "Benchmarked 4+ segmentation models on drone imagery and selected SegFormer-B5 for the cleanest detections, cutting false positives roughly 3x versus baselines.",
                    "Built the OpenCV/NumPy pipeline — tiling, filtering, and per-class NMS — emitting annotated images and structured JSON.",
                ],
            },
        ],
        "projects": [],
        "academic": [],
    },
    "Drone_Robotics_Software_Engineer": {
        "filename": "Resume_Janadri_Yalla_Yashwanth_Drone_Robotics_Software_Engineer",
        "headline": "Drone / Robotics Software Engineer",
        "summary": "Robotics software engineer building autonomous drone systems on ArduPilot — companion-computer vision (Raspberry Pi, AprilTag, OpenCV), MAVLink control, and precision landing on stationary and moving platforms. Also built a from-scratch ArduPilot/DataFlash log parser and a 50+ check flight-diagnostics engine. B.Tech in CSE (AI & ML), CGPA 8.6.",
        "skills": [
            ("Languages", "Python, C"),
            ("Flight / Robotics", "ArduPilot / ArduCopter, MAVLink, pymavlink, Mission Planner, CubeOrange flight controller"),
            ("Companion / Vision", "Raspberry Pi 5, picamera2, AprilTag (pupil-apriltags), OpenCV, camera calibration"),
            ("Control & Systems", "State machines, P-control, velocity feed-forward, GPS/haversine navigation, EKF/IMU/GPS"),
            ("Telemetry / Logs", "ArduPilot DataFlash (.bin) parsing, GPS-week timestamps, sensor diagnostics"),
            ("Practices", "Agile / Scrum, SDLC, code review, version control, testing"),
            ("Tooling", "Linux, systemd, Docker, Git, NumPy"),
        ],
        "tools": "Git & GitHub, JIRA, Claude Code, VS Code, Mission Planner, MAVProxy, Docker, Linux, systemd",
        "projects": ["precland", "flightlog", "aerial", "geoai"],
        "academic": [],
    },
    "Cloud_DevOps_Engineer": {
        "filename": "Resume_Janadri_Yalla_Yashwanth_Cloud_DevOps_Systems_Engineer",
        "headline": "Software & DevOps Engineer | AWS · Docker · CI/CD · Linux · Python | Cloud & Infrastructure Automation",
        "summary": "B.Tech CSE (AI & ML) 2025 graduate, CGPA 8.6, who builds, ships, and operates production systems on Linux — Docker & Docker Compose, GitHub Actions CI/CD, systemd services, AWS S3, and distributed task workers (Celery/Redis). Experienced across the full delivery cycle: development, containerization, cloud deployment, automation, and production troubleshooting (CUDA, DB migrations, GPU/CPU workers). Strong problem-solving, automation, and documentation. Immediate joiner.",
        "skills": [
            ("Cloud / Infrastructure", "AWS (S3, /vsis3 streaming), Docker, Docker Compose, Linux (Ubuntu, Raspberry Pi OS), systemd"),
            ("DevOps / Automation", "GitHub Actions CI/CD, deployment automation, dependency packaging, shell & Python scripting"),
            ("Languages", "Python, Bash, C, SQL, Java"),
            ("Backend / Distributed", "FastAPI, Celery, Redis, PostgreSQL, REST APIs, async workers"),
            ("Systems / Tooling", "Linux administration, Windows, Git, monitoring, debugging & root-cause analysis"),
            ("Fundamentals", "Data Structures & Algorithms, DBMS, Operating Systems, networking basics"),
        ],
        "tools": "Git & GitHub, Claude Code, VS Code, Docker, GitHub Actions, AWS CLI, systemd, Linux",
        "projects": ["survey", "flightlog", "precland"],
        "academic": [],
    },
    "EdgeVerve_Systems_Engineer": {
        "filename": "Resume_Janadri_Yalla_Yashwanth_EdgeVerve_Systems_Engineer",
        "headline": "Software Engineer | B.Tech CSE (AI & ML) 2025 | Python · Java · SQL · Full-Stack Development",
        "summary": "B.Tech in Computer Science (AI & ML) 2025 graduate, CGPA 8.6, with hands-on software development experience across full-stack web applications, Python/FastAPI backends, and AI/ML — currently building and supporting production systems at Marut Drones. Strong CS fundamentals (Data Structures & Algorithms, DBMS, OS) and end-to-end experience across development, testing, and deployment. Immediate joiner seeking a Systems Engineer role to build, customize, and implement enterprise products.",
        "skills": [
            ("Languages", "Python, Java, C, SQL"),
            ("CS Fundamentals", "Data Structures & Algorithms, DBMS, Operating Systems, OOP, REST API design"),
            ("Web / Frameworks", "FastAPI, React, HTML, CSS, Uvicorn (ASGI), async/await"),
            ("Databases", "SQL, PostgreSQL, Redis"),
            ("Tools / DevOps", "Git, Docker, Docker Compose, GitHub Actions CI/CD, Linux"),
            ("AI / ML (bonus)", "PyTorch, OpenCV, TensorFlow, scikit-learn"),
        ],
        "tools": "Git & GitHub, Claude Code, VS Code, Docker, GitHub Actions, Postman, Linux",
        "projects": ["flightlog", "survey", "aerial"],
        "academic": [],
    },
    "Full_Stack_Engineer": {
        "filename": "Resume_Janadri_Yalla_Yashwanth_Full_Stack_Engineer",
        "headline": "Full-Stack Engineer",
        "summary": "Full-stack engineer comfortable across the stack — React + TypeScript front ends and FastAPI / Python back ends, containerized with Docker. Built and solo-shipped a ~7,200-line full-stack drone-diagnostics platform end-to-end, from file-upload UI to binary parsing and PDF report generation. B.Tech in CSE (AI & ML), CGPA 8.6.",
        "skills": [
            ("Languages", "TypeScript, JavaScript, Python, C, SQL"),
            ("Frontend", "React 19, Vite, Tailwind CSS, HTML, CSS, Fetch API, DOM / blob handling"),
            ("Backend", "FastAPI, Uvicorn (ASGI), REST API design, async/await, CORS, Pydantic"),
            ("Data / Infra", "PostgreSQL, Redis, Celery, AWS S3, Azure (learning), SQLAlchemy"),
            ("DevOps", "Docker, Docker Compose, GitHub Actions, Git, Linux"),
            ("ML / CV", "PyTorch, OpenCV, SegFormer, SAM-2 (integrating AI features into apps)"),
        ],
        "tools": "Git & GitHub, Claude Code, VS Code, npm & Vite, Docker, GitHub Actions, Postman, Linux",
        "projects": ["flightlog", "survey", "aerial", "geoai"],
        "academic": [],
    },
}

# ---------- ATS-safe character normalization ----------
# Some older ATS parsers (e.g. legacy Taleo) mangle certain Unicode glyphs.
# Replace the risky ones with plain ASCII equivalents; keep em/en dashes (widely supported).
_REPL = {
    "“": '"', "”": '"', "‘": "'", "’": "'",  # smart quotes
    "×": "x",   # multiplication sign ×  ->  x
    "−": "-",   # minus sign −  ->  hyphen
    "≥": ">=",  # ≥
    "≤": "<=",  # ≤
    "²": "2",   # superscript 2 (m²) -> m2
    "³": "3",   # superscript 3 (m³) -> m3
    "→": "to",  # right arrow
}


def ats(s):
    if not isinstance(s, str):
        return s
    for k, v in _REPL.items():
        s = s.replace(k, v)
    return s


def _sanitize(obj):
    if isinstance(obj, str):
        return ats(obj)
    if isinstance(obj, list):
        return [_sanitize(x) for x in obj]
    if isinstance(obj, tuple):
        return tuple(_sanitize(x) for x in obj)
    if isinstance(obj, dict):
        return {k: _sanitize(v) for k, v in obj.items()}
    return obj


P = _sanitize(P)
ROLES = _sanitize(ROLES)
EDUCATION = _sanitize(EDUCATION)
CERTS = _sanitize(CERTS)
ACADEMIC_ML = _sanitize(ACADEMIC_ML)


# ---------- docx styling helpers ----------
ACCENT = RGBColor(0x1F, 0x3A, 0x5F)   # dark navy
GRAY = RGBColor(0x44, 0x44, 0x44)


def set_margins(doc):
    for s in doc.sections:
        s.top_margin = Inches(0.5)
        s.bottom_margin = Inches(0.5)
        s.left_margin = Inches(0.6)
        s.right_margin = Inches(0.6)


def no_space(p, before=0, after=2, line=1.0):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_bottom_border(p):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3A5F')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_hyperlink(paragraph, url, text, size=9.0, color="0563C1", underline=True, bold=False):
    """Add a clickable external hyperlink run to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rPr.append(rFonts)
    if bold:
        rPr.append(OxmlElement("w:b"))
    col = OxmlElement("w:color")
    col.set(qn("w:val"), color)
    rPr.append(col)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def heading(doc, text):
    p = doc.add_paragraph()
    no_space(p, before=8, after=3)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = ACCENT
    r.font.name = "Calibri"
    add_bottom_border(p)
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    no_space(p, after=2, line=1.0)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size = Pt(10)
    r = p.add_run(text)
    r.font.size = Pt(10)
    return p


def build_resume(role_key, cfg):
    doc = Document()
    set_margins(doc)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    # --- header ---
    p = doc.add_paragraph()
    no_space(p, after=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(NAME)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    no_space(p, after=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(cfg["headline"])
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = GRAY

    p = doc.add_paragraph()
    no_space(p, after=4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep = "  |  "

    def _txt(s):
        rr = p.add_run(s)
        rr.font.size = Pt(9)
        rr.font.color.rgb = GRAY
        return rr

    add_hyperlink(p, "mailto:" + EMAIL, EMAIL)
    _txt(sep)
    _txt(PHONE)
    _txt(sep)
    add_hyperlink(p, LINKEDIN_URL, "LinkedIn")
    _txt(sep)
    add_hyperlink(p, GITHUB_URL, "GitHub")
    _txt(sep)
    _txt(LOCATION)

    # --- summary ---
    heading(doc, "Summary")
    p = doc.add_paragraph()
    no_space(p, after=3)
    r = p.add_run(cfg["summary"])
    r.font.size = Pt(10)

    # --- skills ---
    heading(doc, "Technical Skills")
    skill_rows = list(cfg["skills"])
    if cfg.get("tools"):
        skill_rows.append(("Tools", cfg["tools"]))
    for label, val in skill_rows:
        p = doc.add_paragraph()
        no_space(p, after=1, line=1.0)
        rb = p.add_run(f"{label}:  ")
        rb.bold = True
        rb.font.size = Pt(10)
        r = p.add_run(val)
        r.font.size = Pt(10)

    # --- experience ---
    heading(doc, "Experience")
    p = doc.add_paragraph()
    no_space(p, after=0)
    rt = p.add_run("Software Engineer Trainee — ")
    rt.bold = True
    rt.font.size = Pt(10.5)
    add_hyperlink(p, MARUT_URL, "Marut Drones", size=10.5, color="1F3A5F", underline=False, bold=True)
    rt2 = p.add_run("\tFeb 2026 – Present")
    rt2.font.size = Pt(9.5)
    rt2.italic = True
    # right-align the date with a tab stop
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.2), WD_TAB_ALIGNMENT.RIGHT)
    if cfg.get("marut_bullets"):
        for _b in cfg["marut_bullets"]:
            bullet(doc, _b)
    else:
        bullet(doc, "Software engineer on the drone AI/automation team — shipped production computer-vision detection services, geospatial backend features, and autonomous-flight tooling; selected projects below.")

    p = doc.add_paragraph()
    no_space(p, before=3, after=0)
    rt = p.add_run("Machine Learning Intern — Rooman Technologies")
    rt.bold = True
    rt.font.size = Pt(10.5)
    rt2 = p.add_run("\tJan 2025 – Mar 2025")
    rt2.font.size = Pt(9.5)
    rt2.italic = True
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.2), WD_TAB_ALIGNMENT.RIGHT)
    bullet(doc, cfg.get("rooman_bullet", "Built ML solutions for healthcare; developed and evaluated a sleep-disorder classification model on physiological and behavioral data, achieving 98% accuracy."))

    # --- projects ---
    proj_list = cfg.get("custom_projects") or [P[key] for key in cfg.get("projects", [])]
    if proj_list:
        heading(doc, cfg.get("projects_heading", "Key Projects — Marut Drones"))
    for proj in proj_list:
        p = doc.add_paragraph()
        no_space(p, before=4, after=0)
        rt = p.add_run(proj["title"])
        rt.bold = True
        rt.font.size = Pt(10.5)
        p2 = doc.add_paragraph()
        no_space(p2, after=1)
        rtech = p2.add_run(proj["tech"])
        rtech.italic = True
        rtech.font.size = Pt(9)
        rtech.font.color.rgb = GRAY
        for b in proj["bullets"]:
            bullet(doc, b)

    # --- academic projects (only if listed) ---
    if cfg["academic"]:
        heading(doc, "Additional Projects")
        for ap in ACADEMIC_ML:
            if ap[0] in cfg["academic"]:
                bullet(doc, f"{ap[2]}", bold_prefix=f"{ap[0]} ({ap[1]}) — ")

    # --- education ---
    heading(doc, "Education")
    for school, dates, detail in EDUCATION:
        p = doc.add_paragraph()
        no_space(p, after=0)
        rt = p.add_run(school)
        rt.bold = True
        rt.font.size = Pt(10.5)
        rt2 = p.add_run(f"\t{dates}")
        rt2.font.size = Pt(9.5)
        rt2.italic = True
        p.paragraph_format.tab_stops.add_tab_stop(Inches(7.2), WD_TAB_ALIGNMENT.RIGHT)
        p3 = doc.add_paragraph()
        no_space(p3, after=2)
        r = p3.add_run(detail)
        r.font.size = Pt(10)

    # --- certifications ---
    heading(doc, "Certifications")
    for c in CERTS:
        bullet(doc, c)

    out = os.path.join(OUT_DIR, cfg["filename"] + ".docx")
    doc.save(out)
    return out


if __name__ == "__main__":
    paths = []
    for k, c in ROLES.items():
        paths.append(build_resume(k, c))
    for pth in paths:
        print("WROTE", pth)
